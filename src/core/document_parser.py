import os
import re
import hashlib
from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime

try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

from ..gitnexus.models import AnalysisResult, CodeEntity, EntityType, Relationship

class DocumentParser:
    """文档解析器，支持多种文档格式的结构化知识提取"""
    
    def __init__(self):
        self.supported_formats = {
            '.md': self._parse_markdown,
            '.txt': self._parse_plain_text,
            '.rst': self._parse_plain_text,
            '.html': self._parse_html,
            '.htm': self._parse_html
        }
        
        if PDF_SUPPORT:
            self.supported_formats['.pdf'] = self._parse_pdf
        if DOCX_SUPPORT:
            self.supported_formats['.docx'] = self._parse_docx
            self.supported_formats['.doc'] = self._parse_docx
    
    def detect_format(self, file_path: Path) -> Optional[str]:
        """检测文档格式"""
        ext = file_path.suffix.lower()
        return ext if ext in self.supported_formats else None
    
    def parse_document(self, file_path: Path) -> Optional[AnalysisResult]:
        """解析文档文件，提取结构化知识"""
        ext = self.detect_format(file_path)
        if not ext:
            return None
        
        try:
            parser = self.supported_formats[ext]
            content, entities, relationships = parser(file_path)
            
            file_hash = self._get_file_hash(file_path)
            
            result = AnalysisResult(
                file_path=str(file_path),
                file_hash=file_hash,
                language="zh-CN" if self._is_chinese(content) else "en-US",
                entities=entities,
                relationships=relationships,
                lines_of_code=len(content.splitlines()),
                analyzed_at=datetime.now().isoformat()
            )
            
            return result
            
        except Exception as e:
            print(f"解析文档 {file_path} 失败: {str(e)}")
            return None
    
    def _parse_markdown(self, file_path: Path) -> tuple[str, List[CodeEntity], List[Relationship]]:
        """解析Markdown文档"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        entities = []
        relationships = []
        
        # 提取标题结构
        lines = content.splitlines()
        current_hierarchy = []
        line_num = 0
        
        for line in lines:
            line_num += 1
            heading_match = re.match(r'^(#{1,6})\s+(.+)', line)
            if heading_match:
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                
                # 调整层级结构
                while len(current_hierarchy) >= level:
                    current_hierarchy.pop()
                
                entity_id = self._generate_entity_id(str(file_path), title, line_num)
                parent_id = current_hierarchy[-1] if current_hierarchy else None
                
                entity = CodeEntity(
                    id=entity_id,
                    entity_type=EntityType.MODULE,
                    name=title,
                    file_path=str(file_path),
                    start_line=line_num,
                    end_line=line_num,
                    content=title,
                    parent_id=parent_id,
                    lines_of_code=1
                )
                entities.append(entity)
                
                current_hierarchy.append(entity_id)
        
        # 提取代码块
        code_block_pattern = re.compile(r'```(\w+)?\n([\s\S]*?)\n```', re.MULTILINE)
        for match in code_block_pattern.finditer(content):
            lang = match.group(1) or "unknown"
            code_content = match.group(2)
            start_line = content.count('\n', 0, match.start()) + 1
            end_line = content.count('\n', 0, match.end()) + 1
            
            entity_id = self._generate_entity_id(str(file_path), f"code_block_{start_line}", start_line)
            
            entity = CodeEntity(
                id=entity_id,
                entity_type=EntityType.CODE,
                name=f"代码块 ({lang})",
                file_path=str(file_path),
                start_line=start_line,
                end_line=end_line,
                content=code_content,
                lines_of_code=end_line - start_line + 1
            )
            entities.append(entity)
        
        return content, entities, relationships
    
    def _parse_plain_text(self, file_path: Path) -> tuple[str, List[CodeEntity], List[Relationship]]:
        """解析纯文本文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        entities = []
        relationships = []
        
        # 提取段落
        paragraphs = re.split(r'\n\s*\n', content)
        current_pos = 0
        
        for idx, para in enumerate(paragraphs):
            para = para.strip()
            if not para:
                continue
            
            start_line = content.count('\n', 0, content.find(para, current_pos)) + 1
            end_line = start_line + para.count('\n')
            
            entity_id = self._generate_entity_id(str(file_path), f"paragraph_{idx}", start_line)
            
            entity = CodeEntity(
                id=entity_id,
                entity_type=EntityType.COMMENT,
                name=f"段落 {idx + 1}",
                file_path=str(file_path),
                start_line=start_line,
                end_line=end_line,
                content=para,
                lines_of_code=end_line - start_line + 1
            )
            entities.append(entity)
            
            current_pos = content.find(para, current_pos) + len(para)
        
        return content, entities, relationships
    
    def _parse_html(self, file_path: Path) -> tuple[str, List[CodeEntity], List[Relationship]]:
        """解析HTML文档（简化处理）"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        # 去除HTML标签，提取纯文本
        plain_text = re.sub(r'<[^>]+>', '', content)
        plain_text = re.sub(r'\s+', ' ', plain_text).strip()
        
        entities = []
        relationships = []
        
        entity_id = self._generate_entity_id(str(file_path), "html_content", 1)
        entity = CodeEntity(
            id=entity_id,
            entity_type=EntityType.DOCSTRING,
            name="HTML文档内容",
            file_path=str(file_path),
            start_line=1,
            end_line=len(content.splitlines()),
            content=plain_text,
            lines_of_code=len(content.splitlines())
        )
        entities.append(entity)
        
        return content, entities, relationships
    
    def _parse_pdf(self, file_path: Path) -> tuple[str, List[CodeEntity], List[Relationship]]:
        """解析PDF文档"""
        content = ""
        entities = []
        relationships = []
        
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                page_content = page.extract_text()
                content += f"\n=== 第 {page_num + 1} 页 ===\n{page_content}"
                
                entity_id = self._generate_entity_id(str(file_path), f"page_{page_num + 1}", page_num * 50)
                entity = CodeEntity(
                    id=entity_id,
                    entity_type=EntityType.COMMENT,
                    name=f"PDF 第 {page_num + 1} 页",
                    file_path=str(file_path),
                    start_line=page_num * 50 + 1,
                    end_line=(page_num + 1) * 50,
                    content=page_content,
                    lines_of_code=page_content.count('\n') + 1
                )
                entities.append(entity)
        
        return content, entities, relationships
    
    def _parse_docx(self, file_path: Path) -> tuple[str, List[CodeEntity], List[Relationship]]:
        """解析Word文档"""
        content = ""
        entities = []
        relationships = []
        
        doc = Document(file_path)
        
        for para_num, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if not para_text:
                continue
            
            content += para_text + "\n"
            
            entity_id = self._generate_entity_id(str(file_path), f"paragraph_{para_num}", para_num + 1)
            entity = CodeEntity(
                id=entity_id,
                entity_type=EntityType.COMMENT,
                name=f"段落 {para_num + 1}",
                file_path=str(file_path),
                start_line=para_num + 1,
                end_line=para_num + 1,
                content=para_text,
                lines_of_code=1
            )
            entities.append(entity)
        
        # 处理表格
        for table_num, table in enumerate(doc.tables):
            table_content = ""
            for row in table.rows:
                row_content = [cell.text for cell in row.cells]
                table_content += " | ".join(row_content) + "\n"
            
            entity_id = self._generate_entity_id(str(file_path), f"table_{table_num}", len(entities) + 1)
            entity = CodeEntity(
                id=entity_id,
                entity_type=EntityType.MODULE,
                name=f"表格 {table_num + 1}",
                file_path=str(file_path),
                start_line=len(content.splitlines()) + 1,
                end_line=len(content.splitlines()) + len(table.rows),
                content=table_content,
                lines_of_code=len(table.rows)
            )
            entities.append(entity)
            content += table_content + "\n"
        
        return content, entities, relationships
    
    def _get_file_hash(self, file_path: Path) -> str:
        """计算文件哈希值"""
        hasher = hashlib.sha256()
        with open(file_path, "rb") as f:
            while chunk := f.read(8192):
                hasher.update(chunk)
        return hasher.hexdigest()
    
    def _generate_entity_id(self, file_path: str, name: str, line: int) -> str:
        """生成实体唯一ID"""
        key = f"{file_path}:{name}:{line}"
        return hashlib.md5(key.encode()).hexdigest()[:16]
    
    def _is_chinese(self, text: str) -> bool:
        """判断是否为中文内容"""
        chinese_chars = re.findall(r'[\u4e00-\u9fff]', text)
        return len(chinese_chars) / len(text) > 0.2 if len(text) > 0 else False
